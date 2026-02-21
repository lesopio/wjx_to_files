#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
问卷星公开问卷 -> Word(.docx) 导出脚本

安装依赖:
    pip install requests beautifulsoup4 python-docx

使用:
    python wjx_to_docx.py <问卷链接>
"""

from __future__ import annotations

import re
import sys
import time
import json
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Sequence
from urllib.parse import urlparse

try:
    import requests
    from bs4 import BeautifulSoup, Tag
    from docx import Document
    from docx.shared import Cm
except ImportError as exc:  # pragma: no cover - import guard for runtime
    print(f"[错误] 缺少依赖: {exc}")
    print("请先执行: pip install requests beautifulsoup4 python-docx")
    sys.exit(1)


EXIT_INVALID_ARGS = 1
EXIT_NETWORK = 2
EXIT_PARSE = 3
EXIT_DOCX = 4

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/122.0.0.0 Safari/537.36"
)


@dataclass
class Section:
    name: str
    start_topic: Optional[int] = None
    end_topic: Optional[int] = None


@dataclass
class Question:
    index: int
    topic_id: Optional[int]
    display_no: str
    qtype: str
    required: bool
    stem: str
    options: str
    logic_notes: str
    section: str = "题目列表"


@dataclass
class Survey:
    title: str
    description: str
    source_url: str
    crawl_time: str
    sections: List[Section] = field(default_factory=list)
    questions: List[Question] = field(default_factory=list)


class ExportError(Exception):
    def __init__(self, code: int, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


def log(message: str) -> None:
    print(f"[INFO] {message}")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def tag_text(tag: Optional[Tag]) -> str:
    if not tag:
        return ""
    return normalize_text(tag.get_text(" ", strip=True))


def unique_keep_order(items: Sequence[str]) -> List[str]:
    seen = set()
    result = []
    for item in items:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result


def truncate(text: str, max_len: int = 100) -> str:
    text = normalize_text(text)
    if len(text) <= max_len:
        return text
    return f"{text[: max_len - 3]}..."


def validate_url(url: str) -> bool:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        return False
    host = parsed.netloc.lower()
    allowed = ("wjx.cn", "wjx.top", "sojump.com", "sojump.cn")
    return any(part in host for part in allowed)


def fetch_html(url: str, timeout_sec: int = 20, retries: int = 2) -> str:
    headers = {"User-Agent": USER_AGENT}
    last_error: Optional[Exception] = None

    for attempt in range(retries + 1):
        try:
            response = requests.get(url, headers=headers, timeout=timeout_sec)
            if response.status_code >= 400:
                raise requests.HTTPError(f"HTTP {response.status_code}")
            if response.apparent_encoding:
                response.encoding = response.apparent_encoding
            return response.text
        except requests.RequestException as exc:
            last_error = exc
            if attempt < retries:
                sleep_sec = 2**attempt
                log(f"网络请求失败，{sleep_sec} 秒后重试 ({attempt + 1}/{retries})...")
                time.sleep(sleep_sec)

    raise ExportError(EXIT_NETWORK, f"网络请求失败: {last_error}")


def check_restricted_page(html: str) -> Optional[str]:
    lower_html = html.lower()
    keyword_reasons = {
        "验证码": "检测到验证码页面，当前版本不支持自动处理验证码。",
        "滑块验证": "检测到滑块验证，当前版本不支持自动处理风控验证。",
        "请输入访问密码": "检测到密码访问页面，当前版本不支持密码问卷。",
        "请先登录": "检测到登录限制页面，当前版本不支持登录态问卷。",
        "访问过于频繁": "检测到频率限制，请稍后重试。",
    }
    for keyword, reason in keyword_reasons.items():
        if keyword.lower() in lower_html:
            return reason

    if re.search(r"<input[^>]+type=['\"]password['\"]", html, flags=re.IGNORECASE):
        return "检测到密码输入框，当前版本不支持密码问卷。"

    return None


def parse_topic_id(raw: Optional[str]) -> Optional[int]:
    if not raw:
        return None
    match = re.search(r"(\d+)", str(raw))
    if not match:
        return None
    return int(match.group(1))


def map_qtype(type_code: str) -> str:
    mapping = {
        "1": "填空",
        "3": "单选",
        "4": "多选",
        "5": "量表",
        "6": "矩阵量表",
    }
    return mapping.get(type_code, f"未知类型({type_code or 'N/A'})")


def extract_description(soup: BeautifulSoup) -> str:
    candidates = [
        "#divDesc",
        "#desc",
        ".desc_begin",
        ".description",
        ".survey-desc",
    ]
    for selector in candidates:
        node = soup.select_one(selector)
        text = tag_text(node)
        if text:
            return text

    meta_desc = soup.select_one("meta[property='og:description']")
    if meta_desc and meta_desc.get("content"):
        return normalize_text(str(meta_desc.get("content")))
    return ""


def is_cutfield(node: Tag) -> bool:
    classes = node.get("class", [])
    return "cutfield" in classes


def is_question_field(node: Tag) -> bool:
    classes = node.get("class", [])
    return "field" in classes and "ui-field-contain" in classes


def extract_global_logic_signals(html: str) -> Dict[str, str]:
    signals: Dict[str, str] = {}
    pattern = re.compile(
        r"var\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*([^;]{1,800});", re.IGNORECASE
    )
    for name, value in pattern.findall(html):
        if re.search(r"(rel|jump|logic|skip|display|cond)", name, re.IGNORECASE):
            clean_value = truncate(value, 140)
            signals[name] = clean_value
            continue
        if re.search(r"(rel|jump|logic|skip|display|condition)", value, re.IGNORECASE):
            signals[name] = truncate(value, 140)
    return signals


def extract_option_labels(node: Tag, selector: str) -> List[str]:
    labels: List[str] = []
    for label in node.select(selector):
        text = tag_text(label)
        if text:
            labels.append(text)
    return unique_keep_order(labels)


def extract_scale_options(node: Tag) -> str:
    left = tag_text(node.select_one(".scaleTitle_frist"))
    right = tag_text(node.select_one(".scaleTitle_last"))
    anchors: List[str] = []
    for anchor in node.select(".scale-rating a[val], .scale-div a[val]"):
        val = normalize_text(str(anchor.get("val", "")))
        title = normalize_text(str(anchor.get("title", ""))) or normalize_text(
            str(anchor.get("htitle", ""))
        )
        if not val:
            continue
        anchors.append(f"{val}({title})" if title else val)

    anchors = unique_keep_order(anchors)
    lines: List[str] = []
    if left or right:
        lines.append(f"范围: {left or '-'} -> {right or '-'}")
    if anchors:
        lines.append(f"刻度: {'，'.join(anchors)}")

    return "\n".join(lines) if lines else "（量表题，未提取到刻度文本）"


def extract_matrix_options(node: Tag) -> str:
    table = node.select_one("table.matrix-rating, table.matrixtable")
    if not table:
        return "（矩阵题，未提取到表格结构）"

    rows = table.select("tr")
    if not rows:
        return "（矩阵题，表格内容为空）"

    header_cells = rows[0].find_all(["th", "td"])
    headers = [tag_text(cell) for cell in header_cells]
    if headers and not headers[0]:
        headers = headers[1:]
    headers = [h for h in headers if h]

    row_labels: List[str] = []
    for row in rows[1:]:
        cells = row.find_all(["th", "td"])
        if not cells:
            continue
        first = tag_text(cells[0])
        if first:
            row_labels.append(first)

    lines: List[str] = []
    if headers:
        lines.append(f"列: {' | '.join(headers)}")
    if row_labels:
        lines.append(f"行: {'；'.join(row_labels)}")

    return "\n".join(lines) if lines else "（矩阵题，未提取到有效行列文本）"


def extract_options(node: Tag, type_code: str) -> str:
    if type_code == "3":
        labels = extract_option_labels(node, ".ui-radio .label")
        return "；".join(labels) if labels else "（单选题，未提取到选项）"
    if type_code == "4":
        labels = extract_option_labels(node, ".ui-checkbox .label")
        return "；".join(labels) if labels else "（多选题，未提取到选项）"
    if type_code == "1":
        return "（文本输入）"
    if type_code == "5":
        return extract_scale_options(node)
    if type_code == "6":
        return extract_matrix_options(node)

    fallback_labels = extract_option_labels(node, ".label")
    if fallback_labels:
        return "；".join(fallback_labels)
    return "（无选项或开放题）"


def extract_logic_notes(
    node: Tag, topic_id: Optional[int], global_signals: Dict[str, str]
) -> str:
    notes: List[str] = []

    question_attrs = [
        "relation",
        "jumpto",
        "skipto",
        "rely",
        "cond",
        "condition",
        "showcond",
        "hidecond",
        "display",
    ]
    for attr in question_attrs:
        value = normalize_text(str(node.get(attr, "")))
        if value:
            notes.append(f"题属性 {attr}={value}")

    style = normalize_text(str(node.get("style", ""))).lower().replace(" ", "")
    if "display:none" in style:
        notes.append("题目初始隐藏（style=display:none）")

    for input_tag in node.select("input, option"):
        opt_id = (
            normalize_text(str(input_tag.get("id", "")))
            or normalize_text(str(input_tag.get("name", "")))
            or normalize_text(str(input_tag.get("value", "")))
            or "未知选项"
        )
        jumpto = normalize_text(str(input_tag.get("jumpto", "")))
        rel = normalize_text(str(input_tag.get("rel", "")))
        if jumpto:
            notes.append(f"选项 {opt_id} 跳转到 {jumpto}")
        if rel:
            guess = ""
            match = re.search(r"q(\d+)", rel, flags=re.IGNORECASE)
            if match:
                guess = f"（可能关联题 {match.group(1)}）"
            notes.append(f"选项 {opt_id} 触发关联字段 {rel}{guess}")

    topic_signals: List[str] = []
    if topic_id is not None:
        q_key = f"q{topic_id}".lower()
        for name, value in global_signals.items():
            lower_value = value.lower()
            if q_key in lower_value or f"_{topic_id}" in lower_value:
                topic_signals.append(f"{name}={truncate(value, 70)}")

    notes = unique_keep_order(notes)
    if notes:
        if topic_signals:
            notes.append(f"候选脚本标记: {'；'.join(unique_keep_order(topic_signals[:3]))}")
        return "；".join(notes[:8])

    if topic_signals:
        merged = "；".join(unique_keep_order(topic_signals[:4]))
        return f"存在候选逻辑标记：{merged}（未形成可解析跳题规则）"

    return "未检测到显式逻辑"


def parse_question(
    node: Tag, display_index: int, section_name: str, global_signals: Dict[str, str]
) -> Question:
    topic_id = parse_topic_id(node.get("topic"))
    type_code = normalize_text(str(node.get("type", "")))
    qtype = map_qtype(type_code)

    required_attr = normalize_text(str(node.get("req", ""))) == "1"
    required_star = node.select_one(".field-label .req") is not None
    required = required_attr or required_star

    stem = tag_text(node.select_one(".topichtml")) or "（未识别题干）"
    options = extract_options(node, type_code)
    logic_notes = extract_logic_notes(node, topic_id, global_signals)

    return Question(
        index=display_index,
        topic_id=topic_id,
        display_no=str(display_index),
        qtype=qtype,
        required=required,
        stem=stem,
        options=options,
        logic_notes=logic_notes,
        section=section_name,
    )


def build_sections(questions: List[Question], ordered_names: List[str]) -> List[Section]:
    sections: List[Section] = []
    for name in ordered_names:
        topics = [q.topic_id for q in questions if q.section == name and q.topic_id is not None]
        start_topic = min(topics) if topics else None
        end_topic = max(topics) if topics else None
        if not topics and name != "题目列表":
            continue
        sections.append(Section(name=name, start_topic=start_topic, end_topic=end_topic))

    if not sections:
        sections.append(Section(name="题目列表"))
    if len(sections) > 1:
        sections = [
            s
            for s in sections
            if not (s.name == "题目列表" and s.start_topic is None and s.end_topic is None)
        ]
    return sections


def parse_survey(html: str, source_url: str) -> Survey:
    reason = check_restricted_page(html)
    if reason:
        raise ExportError(EXIT_PARSE, reason)

    soup = BeautifulSoup(html, "html.parser")
    container = soup.select_one("#divQuestion")
    if not container:
        raise ExportError(EXIT_PARSE, "页面中未找到题目容器 #divQuestion。")

    title = tag_text(soup.select_one("#htitle")) or tag_text(soup.title) or "问卷导出"
    description = extract_description(soup)
    crawl_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    global_signals = extract_global_logic_signals(html)

    ordered_names: List[str] = ["题目列表"]
    current_section = "题目列表"
    questions: List[Question] = []

    nodes: List[Tag] = container.find_all("div")
    display_index = 0
    for node in nodes:
        if is_cutfield(node):
            section_name = tag_text(node.find("div")) or tag_text(node)
            if section_name:
                current_section = section_name
                if current_section not in ordered_names:
                    ordered_names.append(current_section)
            continue

        if is_question_field(node):
            display_index += 1
            question = parse_question(node, display_index, current_section, global_signals)
            questions.append(question)

    if not questions:
        raise ExportError(EXIT_PARSE, "页面解析完成，但未识别到题目。")

    sections = build_sections(questions, ordered_names)
    return Survey(
        title=title,
        description=description,
        source_url=source_url,
        crawl_time=crawl_time,
        sections=sections,
        questions=questions,
    )


def sanitize_filename(name: str) -> str:
    name = normalize_text(name)
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name)
    name = name.strip(" .")
    if not name:
        name = "wjx_export"
    if len(name) > 80:
        name = name[:80].rstrip(" .")
    return name


def build_output_paths(title: str) -> Dict[str, Path]:
    base = sanitize_filename(title)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    cwd = Path.cwd()

    for idx in range(0, 1000):
        if idx == 0:
            stem = base
        elif idx == 1:
            stem = f"{base}_{stamp}"
        else:
            stem = f"{base}_{stamp}_{idx}"

        paths = {
            "docx": cwd / f"{stem}.docx",
            "json": cwd / f"{stem}.json",
            "md": cwd / f"{stem}.md",
        }
        if not any(path.exists() for path in paths.values()):
            return paths

    raise ExportError(EXIT_DOCX, "输出文件名冲突过多，请稍后重试。")


def set_table_col_widths(table, widths_cm: Sequence[float]) -> None:
    widths = [Cm(v) for v in widths_cm]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def write_docx(survey: Survey, output_path: Path) -> None:
    try:
        doc = Document()
        doc.add_heading(survey.title, level=0)
        if survey.description:
            doc.add_paragraph(f"简介：{survey.description}")
        doc.add_paragraph(f"来源：{survey.source_url}")
        doc.add_paragraph(f"抓取时间：{survey.crawl_time}")
        doc.add_paragraph(f"题目总数：{len(survey.questions)}")

        for section in survey.sections:
            section_questions = [q for q in survey.questions if q.section == section.name]
            if not section_questions:
                continue
            doc.add_heading(section.name, level=2)

            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            headers = ["题号", "题型", "必填", "题干", "选项", "逻辑"]
            for i, text in enumerate(headers):
                table.rows[0].cells[i].text = text
            set_table_col_widths(table, [1.2, 1.8, 1.0, 4.5, 4.0, 3.0])

            for q in section_questions:
                row = table.add_row().cells
                row[0].text = q.display_no
                row[1].text = q.qtype
                row[2].text = "是" if q.required else "否"
                row[3].text = q.stem
                row[4].text = q.options
                row[5].text = q.logic_notes

            doc.add_paragraph("")

        doc.save(str(output_path))
    except Exception as exc:
        raise ExportError(EXIT_DOCX, f"Word 写入失败: {exc}") from exc


def survey_to_dict(survey: Survey) -> Dict[str, object]:
    question_items: List[Dict[str, object]] = []
    for q in survey.questions:
        question_items.append(
            {
                "index": q.index,
                "topic_id": q.topic_id,
                "display_no": q.display_no,
                "qtype": q.qtype,
                "required": q.required,
                "stem": q.stem,
                "options": q.options,
                "logic_notes": q.logic_notes,
                "section": q.section,
            }
        )

    section_items: List[Dict[str, object]] = []
    for s in survey.sections:
        sec_questions = [q for q in question_items if q["section"] == s.name]
        section_items.append(
            {
                "name": s.name,
                "start_topic": s.start_topic,
                "end_topic": s.end_topic,
                "question_count": len(sec_questions),
                "questions": sec_questions,
            }
        )

    return {
        "schema_version": "1.0",
        "title": survey.title,
        "description": survey.description,
        "source_url": survey.source_url,
        "crawl_time": survey.crawl_time,
        "question_count": len(survey.questions),
        "sections": section_items,
        "questions": question_items,
    }


def write_json(survey: Survey, output_path: Path) -> None:
    try:
        payload = survey_to_dict(survey)
        output_path.write_text(
            json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception as exc:
        raise ExportError(EXIT_DOCX, f"JSON 写入失败: {exc}") from exc


def render_markdown(survey: Survey) -> str:
    lines: List[str] = []
    lines.append(f"# {survey.title}")
    lines.append("")
    lines.append("## 元数据")
    lines.append(f"- 来源: {survey.source_url}")
    lines.append(f"- 抓取时间: {survey.crawl_time}")
    lines.append(f"- 题目总数: {len(survey.questions)}")
    if survey.description:
        lines.append(f"- 简介: {survey.description}")
    lines.append("")

    for section in survey.sections:
        section_questions = [q for q in survey.questions if q.section == section.name]
        if not section_questions:
            continue
        lines.append(f"## 章节: {section.name}")
        lines.append(f"- 起始题号(topic): {section.start_topic if section.start_topic is not None else 'N/A'}")
        lines.append(f"- 结束题号(topic): {section.end_topic if section.end_topic is not None else 'N/A'}")
        lines.append(f"- 章节题数: {len(section_questions)}")
        lines.append("")

        for q in section_questions:
            lines.append(f"### Q{q.display_no}")
            lines.append(f"- index: {q.index}")
            lines.append(f"- topic_id: {q.topic_id if q.topic_id is not None else 'N/A'}")
            lines.append(f"- qtype: {q.qtype}")
            lines.append(f"- required: {'是' if q.required else '否'}")
            lines.append(f"- stem: {q.stem}")
            lines.append("- options:")
            lines.append("```text")
            lines.append(q.options)
            lines.append("```")
            lines.append("- logic_notes:")
            lines.append("```text")
            lines.append(q.logic_notes)
            lines.append("```")
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def write_markdown(survey: Survey, output_path: Path) -> None:
    try:
        output_path.write_text(render_markdown(survey), encoding="utf-8")
    except Exception as exc:
        raise ExportError(EXIT_DOCX, f"Markdown 写入失败: {exc}") from exc


def print_usage() -> None:
    print("用法: python wjx_to_docx.py <问卷链接>")
    print("示例: python wjx_to_docx.py https://v.wjx.cn/vm/wPiiM7G.aspx")


def main(argv: List[str]) -> int:
    if len(argv) != 2:
        print_usage()
        return EXIT_INVALID_ARGS

    url = normalize_text(argv[1])
    if not validate_url(url):
        print("[错误] 链接不合法，或不是支持的问卷星/sojump 域名。")
        print_usage()
        return EXIT_INVALID_ARGS

    try:
        log("开始请求问卷页面...")
        html = fetch_html(url)
        log("开始解析问卷结构...")
        survey = parse_survey(html, url)
        log(f"解析完成: 共识别 {len(survey.questions)} 题。")

        output_paths = build_output_paths(survey.title)
        log("开始生成 Word 文档...")
        write_docx(survey, output_paths["docx"])
        log("开始生成 JSON 文件...")
        write_json(survey, output_paths["json"])
        log("开始生成 Markdown 文件...")
        write_markdown(survey, output_paths["md"])
        log("导出完成:")
        log(f"- DOCX: {output_paths['docx']}")
        log(f"- JSON: {output_paths['json']}")
        log(f"- MD: {output_paths['md']}")
        return 0
    except ExportError as exc:
        print(f"[错误] {exc.message}")
        return exc.code


if __name__ == "__main__":
    sys.exit(main(sys.argv))
